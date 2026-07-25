"""
Documents router — CRUD + semantic search for user documents.
All ingested documents are chunked + embedded and stored via DocumentChunk
for retrieval by DocumentPlugin inside the Cognitive Kernel.

Endpoints:
  GET    /documents              — list user docs
  POST   /documents              — upload + process (multipart)
  GET    /documents/{id}         — document detail
  DELETE /documents/{id}         — delete doc + all chunks
  PUT    /documents/{id}         — re-upload / update title
  POST   /documents/search       — semantic search
"""

from __future__ import annotations

import io
import json
import re
import uuid
import zipfile

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, Request, UploadFile
from pydantic import BaseModel
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from gateway.api_gateway.resource_scope import normalized_tenant_scope, scoped_documents_statement
from gateway.api_gateway.routers.auth import get_current_user
from gateway.api_gateway.tenant_middleware import build_tenant_metadata
from infra.audit.logger import write_audit_log
from infra.config.settings import settings
from infra.errors import AppException, ErrorCodes
from infra.observability.logger import get_logger
from infra.storage.database import db_session_dependency as get_db
from infra.storage.models import Document, KnowledgeSource, Project, User
from knowledge.access import CLASSIFICATION_RANK, require_space_role
from knowledge.jobs import enqueue_document_compile
from plugins.document_plugin import generate_llmwiki_entries
from plugins.document_retrieval import fetch_document_candidates, score_document_candidates
from services import document_ingestion

# 保留旧 Router 私有符号，避免兼容调用方中断。
_has_table = document_ingestion.has_table
_ingest = document_ingestion.ingest_document
_merge_ingest_metadata = document_ingestion.merge_ingest_metadata
_sanitize_text = document_ingestion.sanitize_text

logger = get_logger(__name__)
router = APIRouter()


# ── Schemas ───────────────────────────────────────────────────────────────────


def _source_requires_withdrawal(source: KnowledgeSource) -> bool:
    """已进入企业治理或发布链的来源不能随原始资料被物理删除。"""
    return bool(
        source.space_id
        or source.active_version_id
        or source.status in {"review", "published", "deprecated"}
    )


class DocumentOut(BaseModel):
    id: str
    title: str
    file_type: str
    file_size: int
    chunk_count: int
    chunk_strategy: int = 1
    version: int
    status: str
    project_id: str | None = None
    created_at: str
    updated_at: str
    metadata: dict


class DocumentDetail(DocumentOut):
    content_preview: str


class SearchRequest(BaseModel):
    query: str
    top_k: int = 6
    project_id: str | None = None


class SearchResult(BaseModel):
    document_id: str
    title: str
    chunk_index: int
    content: str
    score: float
    metadata: dict


# ── Text extraction ───────────────────────────────────────────────────────────


def _text_quality_score(text: str) -> float:
    """Simple heuristic: higher means more likely readable natural text."""
    if not text:
        return 0.0
    total = len(text)
    if total == 0:
        return 0.0

    cjk = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    alnum = sum(1 for ch in text if ch.isalnum())
    printable = sum(1 for ch in text if ch.isprintable())
    punct = sum(1 for ch in text if ch in "，。！？；：,.!?;:()（）【】[]《》<>-_")

    # Reward printable/alnum/cjk/punctuation ratios, penalize noise characters.
    score = (
        (printable / total) * 0.45
        + (alnum / total) * 0.30
        + (cjk / total) * 0.20
        + (punct / total) * 0.05
    )
    return round(score, 6)


def _decode_text_best_effort(raw: bytes) -> str:
    # Try common encodings for Chinese and UTF text; choose the best quality result.
    candidates: list[str] = []
    for enc in ("utf-8", "utf-8-sig", "gb18030", "gbk", "big5", "latin1"):
        try:
            candidates.append(raw.decode(enc, errors="strict"))
        except Exception:
            continue

    if not candidates:
        candidates.append(raw.decode("utf-8", errors="replace"))

    best = max(candidates, key=_text_quality_score)
    return _sanitize_text(best)


async def _extract_text(raw: bytes, filename: str) -> str:
    fname = filename.lower()

    if fname.endswith(".pdf"):
        pdf_candidates: list[str] = []

        # 1) pypdf
        try:
            import pypdf  # type: ignore

            reader = pypdf.PdfReader(io.BytesIO(raw))
            txt = "\n".join(page.extract_text() or "" for page in reader.pages)
            pdf_candidates.append(_sanitize_text(txt))
        except Exception as exc:
            logger.debug("pdf_extract_pypdf_skipped", error=str(exc))

        # 2) pymupdf (better for many CJK PDFs)
        try:
            import fitz  # type: ignore

            with fitz.open(stream=raw, filetype="pdf") as doc:
                txt = "\n".join(page.get_text("text") or "" for page in doc)
            pdf_candidates.append(_sanitize_text(txt))
        except Exception as exc:
            logger.debug("pdf_extract_pymupdf_skipped", error=str(exc))

        # Pick highest quality extracted text.
        pdf_candidates = [c for c in pdf_candidates if c]
        if pdf_candidates:
            best_pdf = max(pdf_candidates, key=_text_quality_score)
            return best_pdf

        # PDF binary content should never be decoded as plain UTF-8 text.
        return ""

    if fname.endswith(".docx"):
        # Prefer python-docx; fallback to unzip xml extraction.
        try:
            import docx  # type: ignore

            document = docx.Document(io.BytesIO(raw))
            parts: list[str] = []
            parts.extend((p.text or "") for p in document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    cells = [(cell.text or "").strip() for cell in row.cells]
                    row_text = " ".join(cell for cell in cells if cell)
                    if row_text:
                        parts.append(row_text)
            txt = "\n".join(part for part in parts if part.strip())
            return _sanitize_text(txt)
        except Exception:
            try:
                with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                    with zf.open("word/document.xml") as f:
                        xml = f.read().decode("utf-8", errors="replace")
                xml = re.sub(r"<[^>]+>", " ", xml)
                xml = re.sub(r"\s+", " ", xml)
                return _sanitize_text(xml)
            except Exception:
                return ""

    return _decode_text_best_effort(raw)


# ── Shared ingest service compatibility imports ───────────────────────────────


def _safe_chunk_meta(raw: str | None) -> dict:
    if not raw:
        return {}
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            return data
        return {}
    except Exception:
        return {}


def _doc_out(d: Document) -> DocumentOut:
    try:
        meta = json.loads(d.doc_metadata) if d.doc_metadata else {}
    except Exception:
        meta = {}
    return DocumentOut(
        id=d.id,
        title=d.title,
        file_type=d.file_type,
        file_size=d.file_size,
        chunk_count=d.chunk_count,
        chunk_strategy=getattr(d, "chunk_strategy", 1) or 1,
        version=d.version,
        status=d.status,
        project_id=d.project_id,
        created_at=d.created_at.isoformat(),
        updated_at=d.updated_at.isoformat(),
        metadata=meta,
    )


# ── Routes ────────────────────────────────────────────────────────────────────


@router.get("/documents", response_model=list[DocumentOut])
async def list_documents(
    http_request: Request,
    project_id: str | None = None,
    limit: int = 50,
    offset: int = 0,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> list[DocumentOut]:
    tenant_md = build_tenant_metadata(http_request, user_id=current_user.id)
    result = await db.execute(
        scoped_documents_statement(
            user_id=current_user.id,
            tenant_metadata=tenant_md,
            project_id=project_id,
        )
        .order_by(Document.created_at.desc())
        .limit(limit)
        .offset(offset)
    )
    return [_doc_out(d) for d in result.scalars().all()]


@router.post("/documents", response_model=DocumentOut, status_code=201)
async def upload_document(
    http_request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str | None = Form(None),
    chunk_strategy: int = Form(1),
    project_id: str | None = Form(None),
    knowledge_space_id: str | None = Form(None),
    classification: str = Form("internal"),
    publish_policy: str = Form("auto"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> DocumentOut:
    raw = await file.read()
    filename = file.filename or "document"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    text = await _extract_text(raw, filename)

    tenant_md = build_tenant_metadata(http_request, user_id=current_user.id)
    doc_tenant = str(tenant_md.get("tenant_id") or "default")
    doc_workspace = str(tenant_md.get("workspace_id") or "default")
    if project_id:
        project = await db.scalar(
            select(Project.id).where(
                Project.id == project_id,
                Project.user_id == current_user.id,
                Project.tenant_id == doc_tenant,
                Project.workspace_id == doc_workspace,
                Project.archived_at.is_(None),
            )
        )
        if project is None:
            raise AppException(ErrorCodes.RESOURCE_NOT_FOUND.code, message="Project not found")
    normalized_publish_policy = publish_policy.strip().lower()
    normalized_classification = classification.strip().lower()
    if normalized_classification not in CLASSIFICATION_RANK:
        raise AppException(ErrorCodes.PARAM_INVALID.code, message="Unsupported classification")
    if knowledge_space_id:
        try:
            space, role = await require_space_role(
                db,
                user=current_user,
                tenant_id=doc_tenant,
                workspace_id=doc_workspace,
                space_id=knowledge_space_id,
                required_role="contributor",
            )
        except PermissionError as exc:
            raise AppException(ErrorCodes.PERMISSION_DENIED.code, message=str(exc)) from exc
        normalized_publish_policy = space.publish_policy
        # Contributor 只能保持或提高密级；只有空间管理员可以主动降密。
        if (
            CLASSIFICATION_RANK[normalized_classification]
            < CLASSIFICATION_RANK[space.default_classification]
            and role != "admin"
        ):
            normalized_classification = space.default_classification
    elif normalized_publish_policy not in {"auto", "review"}:
        raise AppException(
            ErrorCodes.PARAM_INVALID.code, message="publish_policy must be auto or review"
        )

    doc = Document(
        id=str(uuid.uuid4()),
        owner_id=current_user.id,
        tenant_id=doc_tenant,
        workspace_id=doc_workspace,
        project_id=project_id,
        title=title or filename,
        file_type=ext,
        file_size=len(raw),
        version=1,
        status="pending",
        chunk_strategy=max(1, min(chunk_strategy, 8)),
        doc_metadata=json.dumps(
            {
                "project_id": project_id,
                "knowledge_space_id": knowledge_space_id,
                "knowledge_steward_id": current_user.id,
                "classification": normalized_classification,
                "source_system": "upload",
                "publish_policy": normalized_publish_policy,
            },
            ensure_ascii=False,
        ),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    try:
        await _ingest(db, doc, text)
    except Exception as exc:
        await db.rollback()
        doc.status = "error"
        safe_msg = str(exc)[:400]
        try:
            meta = json.loads(doc.doc_metadata) if doc.doc_metadata else {}
        except Exception:
            meta = {}
        meta["last_error"] = safe_msg
        doc.doc_metadata = json.dumps(meta, ensure_ascii=False)
        db.add(doc)
        await db.commit()

    await db.refresh(doc)
    await write_audit_log(
        user_id=current_user.id,
        action="document.upload",
        resource_type="document",
        resource_id=doc.id,
        payload={
            "title": doc.title,
            "file_type": doc.file_type,
            "file_size": doc.file_size,
            "status": doc.status,
            "knowledge_space_id": knowledge_space_id,
            "classification": normalized_classification,
        },
    )
    if doc.status == "ready":
        background_tasks.add_task(generate_llmwiki_entries, doc.id)
        if getattr(settings, "knowledge_orchestration_enabled", True) and getattr(
            settings, "knowledge_auto_compile_enabled", True
        ):
            await enqueue_document_compile(doc.id)
    return _doc_out(doc)


@router.get("/documents/{doc_id}", response_model=DocumentDetail)
async def get_document(
    http_request: Request,
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> DocumentDetail:
    tenant_md = build_tenant_metadata(http_request, user_id=current_user.id)
    result = await db.execute(
        scoped_documents_statement(
            user_id=current_user.id,
            tenant_metadata=tenant_md,
            document_id=doc_id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise AppException(ErrorCodes.RESOURCE_NOT_FOUND.code, message="Document not found")
    out = _doc_out(doc)
    return DocumentDetail(**out.model_dump(), content_preview=(doc.content or "")[:500])


@router.delete("/documents/{doc_id}")
async def delete_document(
    http_request: Request,
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict:
    tenant_md = build_tenant_metadata(http_request, user_id=current_user.id)
    tenant_id, workspace_id = normalized_tenant_scope(tenant_md)
    await db.execute(
        text("SELECT pg_advisory_xact_lock(hashtext(:document_id))"),
        {"document_id": doc_id},
    )
    # Use a narrow existence check first so we do not trigger ORM relationship loading.
    result = await db.execute(
        select(Document.id, Document.title).where(
            Document.id == doc_id,
            Document.owner_id == current_user.id,
            Document.tenant_id == tenant_id,
            Document.workspace_id == workspace_id,
        )
    )
    row = result.first()
    if not row:
        raise AppException(ErrorCodes.RESOURCE_NOT_FOUND.code, message="Document not found")

    # 已进入企业知识生命周期的资料必须撤回而不是物理删除，否则会破坏
    # SourceVersion → Page/Claim/Relation → Citation 的审计与引用链。
    if await _has_table(db, "knowledge_sources"):
        governed_sources = list(
            (
                await db.execute(
                    select(KnowledgeSource).where(
                        KnowledgeSource.document_id == doc_id,
                        KnowledgeSource.tenant_id == tenant_id,
                        KnowledgeSource.workspace_id == workspace_id,
                    )
                )
            ).scalars()
        )
        protected_sources = [
            source for source in governed_sources if _source_requires_withdrawal(source)
        ]
        if protected_sources:
            raise AppException(
                ErrorCodes.RESOURCE_EXISTS.code,
                message="该资料已进入企业知识治理，请先在知识治理中心撤回知识来源",
                details={"source_ids": [source.id for source in protected_sources]},
            )
        if governed_sources:
            await db.execute(
                text("DELETE FROM knowledge_sources WHERE document_id = :doc_id"),
                {"doc_id": doc_id},
            )
    # Use raw SQL deletes to avoid ORM cascade / lazy-load behavior entirely.
    if await _has_table(db, "document_llmwiki"):
        await db.execute(
            text("DELETE FROM document_llmwiki WHERE document_id = :doc_id"),
            {"doc_id": doc_id},
        )
    await db.execute(
        text("DELETE FROM document_chunks WHERE document_id = :doc_id"), {"doc_id": doc_id}
    )
    await db.execute(
        text(
            "DELETE FROM documents "
            "WHERE id = :doc_id AND owner_id = :owner_id "
            "AND tenant_id = :tenant_id AND workspace_id = :workspace_id"
        ),
        {
            "doc_id": doc_id,
            "owner_id": current_user.id,
            "tenant_id": tenant_id,
            "workspace_id": workspace_id,
        },
    )
    await db.commit()
    try:
        await write_audit_log(
            user_id=current_user.id,
            action="document.delete",
            resource_type="document",
            resource_id=doc_id,
            payload={"title": row.title},
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("Document delete audit log failed", doc_id=doc_id, error=str(exc))
    return {"deleted": True, "id": doc_id}


@router.put("/documents/{doc_id}", response_model=DocumentOut)
async def update_document(
    http_request: Request,
    doc_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile | None = File(None),
    title: str | None = Form(None),
    chunk_strategy: int | None = Form(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> DocumentOut:
    tenant_md = build_tenant_metadata(http_request, user_id=current_user.id)
    result = await db.execute(
        scoped_documents_statement(
            user_id=current_user.id,
            tenant_metadata=tenant_md,
            document_id=doc_id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise AppException(ErrorCodes.RESOURCE_NOT_FOUND.code, message="Document not found")

    if title:
        doc.title = title
    if chunk_strategy is not None:
        doc.chunk_strategy = max(1, min(chunk_strategy, 8))

    if file:
        raw = await file.read()
        filename = file.filename or "document"
        text = await _extract_text(raw, filename)
        doc.file_size = len(raw)
        doc.version += 1
        await _ingest(db, doc, text)
        if doc.status == "ready":
            background_tasks.add_task(generate_llmwiki_entries, doc.id)
            if getattr(settings, "knowledge_orchestration_enabled", True) and getattr(
                settings, "knowledge_auto_compile_enabled", True
            ):
                await enqueue_document_compile(doc.id)
    else:
        await db.commit()

    await db.refresh(doc)
    return _doc_out(doc)


@router.post("/documents/search", response_model=list[SearchResult])
async def search_documents(
    http_request: Request,
    req: SearchRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> list[SearchResult]:
    tenant_md = build_tenant_metadata(http_request, user_id=current_user.id)
    tenant_id, workspace_id = normalized_tenant_scope(tenant_md)
    candidates = await fetch_document_candidates(
        user_id=current_user.id,
        query=req.query,
        limit=200,
        tenant_id=tenant_id,
        workspace_id=workspace_id,
        project_id=req.project_id,
    )
    scored = await score_document_candidates(query=req.query, candidates=candidates)
    return [
        SearchResult(
            document_id=item.chunk.document_id,
            title=item.title,
            chunk_index=item.chunk.chunk_index,
            content=item.chunk.content,
            score=round(item.score, 4),
            metadata=_safe_chunk_meta(item.chunk.chunk_metadata),
        )
        for item in scored[: req.top_k]
    ]
